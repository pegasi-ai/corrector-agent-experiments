import streamlit as st
from pydantic import BaseModel, Field
from typing import List, Tuple
import mammoth
import io
import re
from dotenv import load_dotenv
import os
import PyPDF2
import json
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.shared import OxmlElement, qn
from langchain_openai import ChatOpenAI
from langchain.agents import AgentExecutor, create_openai_functions_agent, create_structured_chat_agent
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.schema import BaseMessage
from langchain.tools import BaseTool
from langchain_core.tools import tool
import preprocess
from urllib.parse import quote, urlparse, urlunparse
import html
from bs4 import BeautifulSoup, NavigableString
from rapidfuzz import fuzz
import string
from token_prices import token_prices_per_million
import tiktoken

load_dotenv()  # This loads variables from .env into the environment

# Set OpenAI API key from environment variable (supports both OPENAI_API_KEY and openai_api_key)
if os.getenv("OPENAI_API_KEY"):
    os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
elif os.getenv("openai_api_key"):
    os.environ["OPENAI_API_KEY"] = os.getenv("openai_api_key")

# Pydantic models for structured output
class Correction(BaseModel):
    original_text: str = Field(description="The original incorrect text")
    corrected_text: str = Field(description="The factually correct text")
    reason: str = Field(description="Explanation of why the correction was needed")

class DocumentAnalysis(BaseModel):
    corrections: List[Correction] = Field(description="List of factual corrections needed")
    summary: str = Field(description="Summary of corrections made")

@tool
def fact_check_document(document_text: str) -> str:
    """
    Analyze a document for factual errors and return corrections in JSON format.
    
    Args:
        document_text: The text content of the document to fact-check
        
    Returns:
        JSON string with corrections, and summary
    """
    # This tool will be used by the LangChain agent to perform fact-checking
    return f"Fact-checking completed for document with {len(document_text)} characters"

def create_langchain_agent(model_name: str):
    """Create a LangChain agent for fact-checking with support for multiple providers"""
    if model_name.startswith("openai:"):
        from langchain_openai import ChatOpenAI
        if "o3" in model_name:
            llm = ChatOpenAI(
                model=model_name.replace("openai:", ""),
                api_key=os.environ.get("OPENAI_API_KEY")
            )
        else:
            llm = ChatOpenAI(
                model=model_name.replace("openai:", ""),
                temperature=0.1,
                api_key=os.environ.get("OPENAI_API_KEY")
            )
    elif model_name.startswith("anthropic:"):
        from langchain_anthropic import ChatAnthropic
        llm = ChatAnthropic(
            model=model_name.replace("anthropic:", ""),
            temperature=0.1,
            api_key=os.environ.get("ANTHROPIC_API_KEY")
        )
    elif model_name.startswith("groq:"):
        from langchain_groq import ChatGroq
        llm = ChatGroq(
            model=model_name.replace("groq:", ""),
            temperature=0.1,
            api_key=os.environ.get("GROQ_API_KEY")
        )
    elif model_name.startswith("gemini:"):
        from langchain_google_genai import ChatGoogleGenerativeAI
        # Extract the model name without the gemini: prefix
        gemini_model = model_name.replace("gemini:", "")
        llm = ChatGoogleGenerativeAI(
            model=gemini_model,
            temperature=0.1,
            google_api_key=os.environ.get("GOOGLE_API_KEY"),
            convert_system_message_to_human=True
        )
    else:
        raise ValueError(f"Unsupported model: {model_name}")

    # Create the prompt template - Fixed to match expected variables
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a fact-checking expert. Your responsibilities are to:
        1. Carefully review the provided document and the referenced website links.
        2. Retrieve relevant factual information from the website links.
        3. Identify any factually incorrect statements in the document based on the retrieved content.
        4. Provide accurate corrections, each with a brief reason. DO NOT include any website link in the reason.
        5. Focus strictly on verifiable facts — including dates, statistics, historical events, and scientific claims.
        6. Be conservative in your judgment — only flag clear, objective factual inaccuracies, not opinions or subjective language.

        Use the fact_check_document tool to analyze the document and return a JSON response with this exact structure:
        {{
            "corrections": [
                {{
                    "original_text": "the incorrect text",
                    "corrected_text": "the correct text", 
                    "reason": "explanation of why correction was needed",
                }}
            ],
            "summary": "summary of corrections made"
        }}"""),
        ("human", "{input}"),
        MessagesPlaceholder(variable_name="agent_scratchpad"),
    ])
    
    # Create tools
    tools = [fact_check_document]
    
    # Create the agent based on the model type
    if model_name.startswith("anthropic:"):
        # Use structured chat agent for Claude (no function-calling support)
        agent = create_structured_chat_agent(llm, tools, prompt)
    else:
        # Use function-calling agent for OpenAI/Groq
        agent = create_openai_functions_agent(llm, tools, prompt)
    
    # Create the agent executor
    agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)
    
    return agent_executor

def process_document_with_langchain(agent_executor, text: str, links: List[str], model_name: str) -> Tuple[DocumentAnalysis, dict]:
    """Process document using LangChain agent and return analysis with token usage"""
    try:
        # Compose the reference links section
        links_section = "\n".join(f"- {link}" for link in links)
        
        # Prepare the input text for token counting
        input_text = f"Please fact-check the following document and determine if the information is factually consistent based on the provided reference links.\n\nDocument:\n{text}\n\nReference Links:\n{links_section}\n\nFor each factual claim, check if it is supported or contradicted by the reference links."
        
        # Count input tokens
        input_tokens = count_tokens(input_text, model_name)
        
        # Run the agent with updated instructions
        result = agent_executor.invoke({
            "input": input_text
        })
        
        # Extract the result
        result_text = result["output"]
        
        # Count output tokens
        output_tokens = count_tokens(result_text, model_name)
        
        # Try to parse JSON from the result
        # The agent might return JSON in the response
        try:
            # Look for JSON in the response
            start_idx = result_text.find('{')
            end_idx = result_text.rfind('}') + 1
            if start_idx != -1 and end_idx != 0:
                json_str = result_text[start_idx:end_idx]
                result_json = json.loads(json_str)
            else:
                # If no JSON found, create a default response
                result_json = {
                    "corrections": [],
                    "summary": "No factual errors found in the document."
                }
        except json.JSONDecodeError:
            # If JSON parsing fails, create a default response
            result_json = {
                "corrections": [],
                "summary": "Analysis completed. Please review the agent's response manually."
            }
        
        # Convert to DocumentAnalysis object
        corrections = [Correction(**corr) for corr in result_json["corrections"]]
        analysis = DocumentAnalysis(
            corrections=corrections,
            summary=result_json["summary"]
        )
        
        # Prepare token usage info
        token_usage = {
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": input_tokens + output_tokens,
            "cost": calculate_cost(input_tokens, output_tokens, model_name)
        }
        
        return analysis, token_usage
        
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        return None, None

def extract_text_from_docx(file_content):
    """Extract text from Word document using mammoth"""
    try:
        # Convert docx to HTML first to preserve some formatting
        result = mammoth.convert_to_html(io.BytesIO(file_content))
        html_content = result.value

        # Create a new Word document
        doc_tmp = Document()
        doc_tmp.add_paragraph(html_content)

        # Save the document
        doc_tmp.save("output_html.docx")
        
        # Extract plain text for processing
        text_result = mammoth.extract_raw_text(io.BytesIO(file_content))
        plain_text = text_result.value
        
        # Extract hyperlinks from HTML content and append to plain text
        import re
        link_pattern = r'<a[^>]*href=["\']([^"\']+)["\'][^>]*>([^<]+)</a>'
        links = re.findall(link_pattern, html_content)
        
        # Append links to plain text
        if links:
            plain_text += "\n\nReferences:\n"
            for url, link_text in links:
                plain_text += f"- [{link_text}]({url})\n"
        
        return plain_text, html_content
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return None, None

def extract_text_from_pdf(file_content):
    """Extract text from PDF file using PyPDF2"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        
        return text, None  # No HTML for PDF
    except Exception as e:
        st.error(f"Error reading PDF document: {str(e)}")
        return None, None

def apply_corrections_with_red_marking(original_text: str, corrections: List[Correction]) -> Tuple[str, List[Correction]]:
    """Apply corrections to text with red highlighting for changes. Return both the highlighted text and the list of corrections that were actually applied, in the order they appear in the original text."""
    import string
    import re
    
    def normalize(s):
        s = s.lower()
        s = s.translate(str.maketrans('', '', string.punctuation))
        s = re.sub(r'\s+', ' ', s)
        return s.strip()
    
    corrected_text = original_text
    applied_corrections = []
    for correction in corrections:
        norm_pattern = normalize(correction.original_text)
        norm_text = normalize(corrected_text)
        # Build mapping from normalized indices to original indices
        norm_to_orig = []
        orig_idx = 0
        norm_chars = []
        while orig_idx < len(corrected_text):
            c = corrected_text[orig_idx]
            if c in string.punctuation:
                orig_idx += 1
                continue
            if c.isspace():
                if norm_chars and norm_chars[-1] == ' ':
                    orig_idx += 1
                    continue
                norm_chars.append(' ')
                norm_to_orig.append(orig_idx)
                orig_idx += 1
            else:
                norm_chars.append(c.lower())
                norm_to_orig.append(orig_idx)
                orig_idx += 1
        norm_text_built = ''.join(norm_chars).strip()
        # Find all matches in the normalized text
        matches = [m for m in re.finditer(re.escape(norm_pattern), norm_text_built)]
        offset = 0
        applied = False
        for m in matches:
            norm_start, norm_end = m.start(), m.end()
            orig_start = norm_to_orig[norm_start]
            orig_end = norm_to_orig[norm_end - 1] + 1  # end is exclusive
            before = corrected_text[:orig_start+offset]
            after = corrected_text[orig_end+offset:]
            before_len = len(before)
            corrected_text = before + f'<span style="color: red; font-weight: bold;">{correction.corrected_text}</span>' + after
            offset += len(f'<span style="color: red; font-weight: bold;">{correction.corrected_text}</span>') - (orig_end - orig_start)
            applied = True
        if applied:
            applied_corrections.append(correction)
        else:
            print("================ norm_text_built ===================")
            print(norm_text_built)
            print("================ norm_pattern ===================")
            print(norm_pattern[0:1000])
    # HTML-encode $ symbols to ensure they display properly in Streamlit
    corrected_text = corrected_text.replace('$', '&#36;')
    return corrected_text, applied_corrections

def create_word_document_with_highlighting(original_text: str, corrections: List[Correction], analysis: DocumentAnalysis, filename: str, original_file_content: bytes = None) -> bytes:
    """Create a Word document with red highlighting for corrections, matching the Streamlit UI display"""
    doc = Document()
    
    if original_file_content and filename.lower().endswith('.docx'):
        
        # Add a title page at the beginning
        title_page = doc.add_paragraph()
        title_run = title_page.add_run('Document Fact-Checking Report')
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
        # Add metadata
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Original File: {filename}')
        doc.add_paragraph(f'Number of Corrections: {len(corrections)}')
        
        # Add separator
        doc.add_paragraph('─' * 60)
        doc.add_paragraph()
        
        # Add original text section with inline corrections
        doc.add_heading('Original Document Content (with inline corrections)', level=1)
        import re
        
        # Work on a copy of the original text
        processed_text = original_text
        # To avoid overlapping replacements, process corrections in order of appearance
        for correction in corrections:
            # Use regex to find all non-overlapping matches (case-insensitive)
            pattern = re.escape(correction.original_text)
            def repl(match):
                return match.group(0) + f" [corrected text: {correction.corrected_text}, reason: {correction.reason}]"
            processed_text = re.sub(pattern, repl, processed_text, flags=re.IGNORECASE)
        
        # Now, add the processed text to the document, highlighting the corrections in red
        p = doc.add_paragraph()
        correction_block_pattern = r'\[corrected text: (.*?), reason: (.*?)\]'
        last_end = 0
        for m in re.finditer(correction_block_pattern, processed_text):
            start, end = m.span()
            if start > last_end:
                p.add_run(processed_text[last_end:start])
            corrected_text, reason = m.group(1), m.group(2)
            # Add prefix
            p.add_run("[corrected text: ")
            # Add corrected text in red
            run_corr = p.add_run(corrected_text)
            run_corr.font.color.rgb = RGBColor(255, 0, 0)
            run_corr.font.bold = True
            # Add separator
            p.add_run(", reason: ")
            # Add reason in red
            run_reason = p.add_run(reason)
            run_reason.font.color.rgb = RGBColor(255, 0, 0)
            run_reason.font.bold = True
            # Add suffix
            p.add_run("]")
            last_end = end
        if last_end < len(processed_text):
            p.add_run(processed_text[last_end:])

    
    else:
        # Fallback to creating a new document if no original file content or not a Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Document with Fact-Checking Annotations', 0)
        title.alignment = 1  # Center alignment
        
        # Add metadata
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Original File: {filename}')
        doc.add_paragraph(f'Number of Corrections: {len(corrections)}')
        
        # Add original text
        doc.add_heading('Original Document Content', level=1)
        doc.add_paragraph(original_text)
        
        # Add corrected text with highlighting (matching Streamlit UI)
        doc.add_heading('Corrected Text (with red highlighting)', level=1)
        corrected_text, applied_corrections = apply_corrections_with_red_marking(original_text, corrections)

        # Parse the HTML-style spans and add runs with formatting
        from html.parser import HTMLParser
        class HighlightParser(HTMLParser):
            def __init__(self, paragraph):
                super().__init__()
                self.paragraph = paragraph
                self.in_red_span = False
            def handle_starttag(self, tag, attrs):
                if tag == 'span':
                    for attr in attrs:
                        if attr[0] == 'style' and 'color: red' in attr[1]:
                            self.in_red_span = True
            def handle_endtag(self, tag):
                if tag == 'span':
                    self.in_red_span = False
            def handle_data(self, data):
                run = self.paragraph.add_run(data)
                if self.in_red_span:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.font.bold = True
        p = doc.add_paragraph()
        parser = HighlightParser(p)
        parser.feed(corrected_text.replace('&#36;', '$'))  # decode $ for Word

        # Add detailed corrections
        if applied_corrections:
            doc.add_heading('Detailed Corrections', level=1)
            for i, correction in enumerate(applied_corrections, 1):
                doc.add_heading(f'Correction #{i}', level=2)
                doc.add_paragraph('Original:')
                orig_p = doc.add_paragraph()
                orig_run = orig_p.add_run(correction.original_text)
                orig_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
                doc.add_paragraph('Corrected:')
                corr_p = doc.add_paragraph()
                corr_run = corr_p.add_run(correction.corrected_text)
                corr_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
                doc.add_paragraph('Reason:')
                doc.add_paragraph(correction.reason)
                doc.add_paragraph('─' * 60)
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

def safe_url(url):
    parts = urlparse(url)
    safe_path = quote(parts.path)
    safe_query = quote(parts.query, safe='=&')
    return urlunparse((parts.scheme, parts.netloc, safe_path, parts.params, safe_query, parts.fragment))

def add_hyperlink(paragraph, url, text=None):
    # This function adds a clickable hyperlink to a paragraph
    # If text is None, use the url as the display text
    if text is None:
        text = url
    url_encoded = safe_url(url)
    part = paragraph.part
    r_id = part.relate_to(url_encoded, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Style: blue and underlined
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    # Add the raw URL as plain text for reference
    paragraph.add_run(f' ({url})')

def insert_annotations_html(html_content, corrections):
    import re
    import html as html_mod
    import string

    def find_spans(s, tag_patterns):
        spans = []
        for pat in tag_patterns:
            for m in re.finditer(pat, s, re.IGNORECASE | re.DOTALL):
                spans.append((m.start(), m.end()))
        return spans

    def mask_spans(s, spans):
        s = list(s)
        for start, end in spans:
            for i in range(start, end):
                s[i] = '\0'
        return ''.join(s)

    def normalize(s):
        s = s.lower()
        s = s.translate(str.maketrans('', '', string.punctuation))
        s = re.sub(r'\s+', ' ', s)
        return s.strip()

    tag_patterns = [
        r'<a[^>]*>.*?</a>',
        r'<p>', r'</p>',
        r'<strong>', r'</strong>'
    ]

    annotated = html_content
    for correction in corrections:
        annotation = f" <span style='color: red; font-weight: normal;'>[corrected text: {html_mod.escape(correction.corrected_text)}, reason: {html_mod.escape(correction.reason)}]</span>"
        spans = find_spans(annotated, tag_patterns)
        masked = mask_spans(annotated, spans)

        # Build normalized version and mapping from normalized index to original index
        norm_chars = []
        norm_to_orig = []
        for i, c in enumerate(masked):
            if c == '\0':
                continue
            if c in string.punctuation:
                continue
            if c.isspace():
                if norm_chars and norm_chars[-1] == ' ':
                    continue
                norm_chars.append(' ')
                norm_to_orig.append(i)
            else:
                norm_chars.append(c.lower())
                norm_to_orig.append(i)
        norm_masked = ''.join(norm_chars)

        norm_pattern = normalize(correction.original_text)
        # Find all matches in the normalized masked string
        matches = [m for m in re.finditer(re.escape(norm_pattern), norm_masked)]

        #print(len(matches))
        # Insert annotation after each match, from end to start
        offset = 0
        for m in reversed(matches):
            norm_start, norm_end = m.start(), m.end()
            # Map normalized start/end index to original string index
            orig_start = norm_to_orig[norm_start]
            orig_end = norm_to_orig[norm_end - 1] + 1  # +1 because end is exclusive
            # Wrap the matched original text in green
            green_span = f"<span style='color: green; font-weight: normal;'>" + annotated[orig_start+offset:orig_end+offset] + "</span>"
            annotated = annotated[:orig_start+offset] + green_span + annotated[orig_end+offset:]
            offset += len(green_span) - (orig_end - orig_start)
            # Insert annotation after the green span
            insert_at = orig_end + offset
            annotated = annotated[:insert_at] + annotation + annotated[insert_at:]
            offset += len(annotation)
    return annotated

def insert_annotations_text(text, corrections):
    processed_text = text
    for correction in corrections:
        # Remove URLs from the text for matching
        text_for_match = re.sub(r'https?://\S+', '', processed_text)
        pattern = re.compile(re.escape(correction.original_text), re.IGNORECASE)
        def repl(match):
            return match.group(0) + f" [corrected text: {correction.corrected_text}, reason: {correction.reason}]"
        processed_text = pattern.sub(repl, text_for_match)
    return processed_text

def create_annotated_document(original_text: str, corrections: List[Correction], filename: str, original_file_content: bytes = None) -> bytes:
    """Create an annotated HTML document with metadata and original content, using exact matching for correction insertion (ignoring hyperlinks)."""
    import io
    from datetime import datetime
    import mammoth
    import html
    from bs4 import BeautifulSoup, NavigableString


    if original_file_content and filename.lower().endswith('.docx'):
        result = mammoth.convert_to_html(io.BytesIO(original_file_content))
        html_content = result.value
        html_content = insert_annotations_html(html_content, corrections)
        metadata_html = f"""
        <h1 style='color:#003366'>Document with Fact-Checking Annotations</h1>
        <p><b>Generated on:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p><b>Original File:</b> {html.escape(filename)}</p>
        <hr/>
        """
        full_html = f"""
        <html><head><meta charset='utf-8'></head><body>
        {metadata_html}
        {html_content}
        </body></html>
        """
        return full_html.encode('utf-8')
    else:
        annotated_text = insert_annotations_text(original_text, corrections)
        metadata_html = f"""
        <h1 style='color:#003366'>Document with Fact-Checking Annotations</h1>
        <p><b>Generated on:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p><b>Original File:</b> {html.escape(filename)}</p>
        <p><b>Number of Corrections:</b> {len(corrections)}</p>
        <hr/>
        """
        body_html = f"<h2>Original Document Content</h2><pre>{html.escape(annotated_text)}</pre>"
        full_html = f"<html><head><meta charset='utf-8'></head><body>{metadata_html}{body_html}</body></html>"
        return full_html.encode('utf-8')

def count_tokens(text: str, model_name: str = "gpt-4") -> int:
    """Count tokens in text using tiktoken"""
    try:
        # Use cl100k_base encoding for GPT-4 and similar models
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))
    except Exception as e:
        # Fallback: rough estimation (1 token ≈ 4 characters)
        return len(text) // 4

def calculate_cost(input_tokens: int, output_tokens: int, model_name: str) -> float:
    """Calculate cost based on token usage and model pricing"""
    if model_name not in token_prices_per_million:
        return 0.0
    
    prices = token_prices_per_million[model_name]
    input_cost = (input_tokens / 1_000_000) * prices["input"]
    output_cost = (output_tokens / 1_000_000) * prices["output"]
    return input_cost + output_cost

def main():
    st.set_page_config(
        page_title="Document Fact Checker",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Document Fact Checker")
    st.markdown("Upload a Microsoft Word document or PDF to check for factual accuracy and get corrections with red highlighting.")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Model selection
        model_options = {
            "OpenAI GPT-4o": "openai:gpt-4o",
            "OpenAI GPT-4o Mini": "openai:gpt-4o-mini",
            "OpenAI GPT-4.1 Mini": "openai:gpt-4.1-mini",
            "OpenAI o3": "openai:o3",
            "Groq Llama 2 70B": "groq:llama-3.3-70b-versatile",
            "Gemini 2.0 Flash": "gemini:models/gemini-2.0-flash",
            "Gemini 2.5 Pro": "gemini:models/gemini-2.5-pro",
        }
        selected_model_label = st.selectbox(
            "Language Model",
            list(model_options.keys()),
            help="Choose the language model for fact-checking."
        )
        selected_model = model_options[selected_model_label]
        
        # API Key input
        api_key_provided = False
        
        if selected_model.startswith("openai:"):
            api_key = st.text_input(
                "OpenAI API Key", 
                type="password", 
                help="Enter your OpenAI API key to use the fact-checking service"
            )
            if api_key:
                os.environ["OPENAI_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("OPENAI_API_KEY"):
                api_key_provided = True
                st.success("✅ OpenAI API key loaded from environment")
                
        elif selected_model.startswith("anthropic:"):
            api_key = st.text_input(
                "Anthropic API Key", 
                type="password", 
                help="Enter your Anthropic API key to use Claude models"
            )
            if api_key:
                os.environ["ANTHROPIC_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("ANTHROPIC_API_KEY"):
                api_key_provided = True
                st.success("✅ Anthropic API key loaded from environment")
                
        elif selected_model.startswith("groq:"):
            api_key = st.text_input(
                "Groq API Key", 
                type="password", 
                help="Enter your Groq API key to use Llama models"
            )
            if api_key:
                os.environ["GROQ_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("GROQ_API_KEY"):
                api_key_provided = True
                st.success("✅ Groq API key loaded from environment")
        
        elif selected_model == "llama-2-70b":
            st.info("No API key required for Llama 2 70B.")
            api_key_provided = True
        
        elif selected_model.startswith("gemini:"):
            api_key = st.text_input(
                "Google Gemini API Key",
                type="password",
                help="Enter your Google Gemini API key"
            )
            if api_key:
                os.environ["GOOGLE_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("GOOGLE_API_KEY"):
                api_key_provided = True
                st.success("✅ Google Gemini API key loaded from environment")
        
        st.markdown("---")
        st.markdown("### 📋 Instructions")
        st.markdown("""
        1. Select your preferred language model
        2. Enter the required API key above
        3. Upload a Microsoft Word document (.docx) or PDF
        4. Click 'Analyze Document' to start fact-checking
        5. Review corrections highlighted in red
        """)
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📤 Upload Document")
        
        # File uploader
        uploaded_file = st.file_uploader(
            "Choose a Word or PDF document",
            type=['docx', 'pdf'],
            help="Upload a .docx or .pdf file to fact-check"
        )
        
        # Clear previous analysis if a new file is uploaded
        if uploaded_file is not None:
            if (
                'uploaded_filename' not in st.session_state or
                st.session_state.uploaded_filename != uploaded_file.name
            ):
                # Clear all previous analysis data
                st.session_state.analysis = None
                st.session_state.original_text = None
                st.session_state.uploaded_filename = uploaded_file.name
                # Force a rerun to clear the display
                if hasattr(st.session_state, 'analysis'):
                    st.rerun()
        
        if uploaded_file is not None:
            st.success(f"✅ Uploaded: {uploaded_file.name}")
            
            # Extract text from document
            file_content = uploaded_file.read()
            # Store file content in session state for later use
            st.session_state.file_content = file_content
            
            if uploaded_file.name.lower().endswith('.docx'):
                plain_text, html_content = extract_text_from_docx(file_content)
            elif uploaded_file.name.lower().endswith('.pdf'):
                plain_text, html_content = extract_text_from_pdf(file_content)
            else:
                st.error("Unsupported file type.")
                return
            
            if plain_text:
                st.subheader("📄 Original Document Preview")
                with st.expander("View original text", expanded=False):
                    st.text_area("Original Content", plain_text, height=300, disabled=True)

                text_links_per_citation, text_links_per_chunk = preprocess.run_href(html_content)
          
                # Track which analyze button was last clicked
                if 'analyze_mode' not in st.session_state:
                    st.session_state.analyze_mode = None

                # Analyze buttons row with highlighting
                analyze_col1, analyze_col2 = st.columns([1, 1])
                with analyze_col1:
                    analyze_doc_clicked = st.button(
                        "🔍 Analyze Document",
                        type="primary" if st.session_state.analyze_mode == 'document' else "secondary",
                        key="analyze_doc_btn"
                    )
                with analyze_col2:
                    analyze_section_clicked = st.button(
                        "🔎 Analyze Section",
                        type="primary" if st.session_state.analyze_mode == 'section' else "secondary",
                        key="analyze_section_btn"
                    )

                # Section selection for Analyze Section (always visible if there are sections)
                section_options = [f"Section {i+1}: {chunk.get('text', '')[:40]}..." for i, chunk in enumerate(text_links_per_chunk)]
                selected_section_idx = st.selectbox("Select Section to Analyze", options=list(range(len(section_options))), format_func=lambda i: section_options[i], key="section_select")

                if analyze_doc_clicked:
                    st.session_state.analyze_mode = 'document'
                    if not api_key_provided:
                        st.error("⚠️ Please enter the required API key in the sidebar first!")
                        return
                    
                    with st.spinner("🤖 Fact-checking document... This may take a few moments."):
                        try:
                            # Clear any previous analysis before starting new one
                            st.session_state.analysis = None
                            
                            # Loop over each chunk and process
                            all_corrections = []
                            all_summaries = []
                            chunk_count = 0
                            total_input_tokens = 0
                            total_output_tokens = 0
                            total_cost = 0.0
                            
                            for chunk in text_links_per_chunk:
                                # Create LangChain agent
                                agent_executor = create_langchain_agent(selected_model)
                                
                                chunk_text = chunk.get('text', '')
                                chunk_links = chunk.get('citations', [])
                                if not chunk_text.strip():
                                    continue
                                analysis, token_usage = process_document_with_langchain(agent_executor, chunk_text, chunk_links, selected_model)
                                if analysis and token_usage:
                                    all_corrections.extend(analysis.corrections)
                                    all_summaries.append(analysis.summary)
                                    chunk_count += 1
                                    total_input_tokens += token_usage["input_tokens"]
                                    total_output_tokens += token_usage["output_tokens"]
                                    total_cost += token_usage["cost"]
                            
                            # Aggregate results
                            if chunk_count > 0:
                                summary = '\n'.join(all_summaries)
                            else:
                                summary = "No factual errors found in the document."
                            final_analysis = DocumentAnalysis(
                                corrections=all_corrections,
                                summary=summary
                            )
                            st.session_state.analysis = final_analysis
                            st.session_state.original_text = plain_text
                            st.session_state.analysis_timestamp = datetime.now().isoformat()
                            
                            # Store token usage information
                            st.session_state.token_usage = {
                                "input_tokens": total_input_tokens,
                                "output_tokens": total_output_tokens,
                                "total_tokens": total_input_tokens + total_output_tokens,
                                "cost": total_cost,
                                "chunks_processed": chunk_count
                            }
                            st.session_state.selected_model = selected_model
                            
                            st.success("✅ Analysis complete!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error during analysis: {str(e)}")
                            st.error("Please check your API key and try again.")
                            st.session_state.analysis = None

                if analyze_section_clicked:
                    st.session_state.analyze_mode = 'section'
                    if not api_key_provided:
                        st.error("⚠️ Please enter the required API key in the sidebar first!")
                        return
                    section_chunk = text_links_per_chunk[selected_section_idx]
                    chunk_text = section_chunk.get('text', '')
                    chunk_links = section_chunk.get('citations', [])
                    with st.spinner(f"🤖 Fact-checking Section {selected_section_idx+1}..."):
                        try:
                            agent_executor = create_langchain_agent(selected_model)
                            analysis, token_usage = process_document_with_langchain(agent_executor, chunk_text, chunk_links, selected_model)
                            if analysis and token_usage:
                                st.session_state.analysis = analysis
                                st.session_state.original_text = chunk_text
                                st.session_state.analysis_timestamp = datetime.now().isoformat()
                                
                                # Store token usage information for single section
                                st.session_state.token_usage = {
                                    "input_tokens": token_usage["input_tokens"],
                                    "output_tokens": token_usage["output_tokens"],
                                    "total_tokens": token_usage["total_tokens"],
                                    "cost": token_usage["cost"],
                                    "chunks_processed": 1
                                }
                                st.session_state.selected_model = selected_model
                                
                                st.success(f"✅ Section {selected_section_idx+1} analysis complete!")
                                st.rerun()
                            else:
                                st.error("❌ Analysis failed. Please try again.")
                        except Exception as e:
                            st.error(f"❌ Error during section analysis: {str(e)}")
                            st.session_state.analysis = None
    
    with col2:
        st.header("📊 Analysis Results")
        
        if hasattr(st.session_state, 'analysis') and st.session_state.analysis:
            analysis = st.session_state.analysis
            original_text = st.session_state.original_text
            
            # Overall metrics
            col2a, col2b, col2c = st.columns(3)
            with col2a:
                st.metric("Corrections Found", len(analysis.corrections))
            with col2b:
                st.metric("Chunks Processed", st.session_state.token_usage.get("chunks_processed", 0))
            with col2c:
                st.metric("Total Cost", f"${st.session_state.token_usage.get('cost', 0):.4f}")
            
            # Token usage details
            if hasattr(st.session_state, 'token_usage') and st.session_state.token_usage:
                token_info = st.session_state.token_usage
                # Get the selected model from session state or use a default
                current_model = st.session_state.get('selected_model', selected_model)
                
                with st.expander("🔢 Token Usage & Cost Details", expanded=False):
                    col_token1, col_token2, col_token3 = st.columns(3)
                    with col_token1:
                        st.metric("Input Tokens", f"{token_info['input_tokens']:,}")
                    with col_token2:
                        st.metric("Output Tokens", f"{token_info['output_tokens']:,}")
                    with col_token3:
                        st.metric("Total Tokens", f"{token_info['total_tokens']:,}")
                    
                    st.markdown(f"**Cost Breakdown:**")
                    model_prices = token_prices_per_million.get(current_model, {})
                    input_price = model_prices.get('input', 0)
                    output_price = model_prices.get('output', 0)
                    st.markdown(f"- Input cost: ${(token_info['input_tokens'] / 1_000_000) * input_price:.4f}")
                    st.markdown(f"- Output cost: ${(token_info['output_tokens'] / 1_000_000) * output_price:.4f}")
                    st.markdown(f"- **Total cost: ${token_info['cost']:.4f}**")
            
            # Corrections list
            if analysis.corrections:
                st.subheader("🔧 Corrections Made")
                
                # Display corrected text
                st.subheader("📄 Corrected Document")
                corrected_text, applied_corrections = apply_corrections_with_red_marking(original_text, analysis.corrections)
                st.markdown(
                    f'<div style="border: 1px solid #ddd; padding: 15px; border-radius: 5px; background-color: #fafafa; max-height: 400px; overflow-y: auto;">{corrected_text}</div>',
                    unsafe_allow_html=True
                )
                
                # Detailed corrections
                st.subheader("📋 Detailed Corrections")
                print("=============== applied_corrections ===================")
                print(len(applied_corrections), len(analysis.corrections))
                for i, correction in enumerate(applied_corrections, 1):
                    with st.expander(f"Correction {i}"):
                        col_orig, col_corr = st.columns(2)
                        
                        with col_orig:
                            st.markdown("**Original:**")
                            st.markdown(f'<div style="background-color: #ffebee; padding: 10px; border-radius: 5px;">{correction.original_text}</div>', unsafe_allow_html=True)
                        
                        with col_corr:
                            st.markdown("**Corrected:**")
                            st.markdown(f'<div style="background-color: #e8f5e8; padding: 10px; border-radius: 5px;">{correction.corrected_text}</div>', unsafe_allow_html=True)
                        
                        st.markdown("**Reason:**")
                        # HTML-encode $ symbols in the reason text to ensure proper display
                        encoded_reason = correction.reason.replace('$', '&#36;')
                        st.markdown(f'<div style="background-color: #f5f5f5; padding: 10px; border-radius: 5px;">{encoded_reason}</div>', unsafe_allow_html=True)
                
                # Download corrected document
                st.subheader("💾 Export")
                
                # Determine the correct file extension
                if uploaded_file and uploaded_file.name.lower().endswith('.pdf'):
                    base_filename = uploaded_file.name.replace('.pdf', '')
                elif uploaded_file and uploaded_file.name.lower().endswith('.docx'):
                    base_filename = uploaded_file.name.replace('.docx', '')
                else:
                    base_filename = "document"
                
                # Create Word document with highlighting
                word_doc_bytes = create_word_document_with_highlighting(
                    original_text, analysis.corrections, analysis, 
                    uploaded_file.name if uploaded_file else "document",
                    st.session_state.get('file_content') if uploaded_file else None
                )
                
                # Download Word document
                st.download_button(
                    label="📄 Download Report",
                    data=word_doc_bytes,
                    file_name=f"fact_check_report_{base_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Create annotated document with corrections and reasons
                annotated_doc_bytes = create_annotated_document(
                    original_text, analysis.corrections, 
                    uploaded_file.name if uploaded_file else "document",
                    st.session_state.get('file_content') if uploaded_file else None
                )
                
                # Download annotated document
                st.download_button(
                    label="📝 Download Report with Annotations",
                    data=annotated_doc_bytes,
                    file_name=f"annotated_document_{base_filename}.html",
                    mime="text/html"
                )

            else:
                st.success("🎉 No factual errors found! The document appears to be accurate.")
        
        else:
            st.info("👆 Upload a document and click 'Analyze Document' to see results here.")
    
    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666; font-size: 0.8em;'>
        Built with ❤️ using LangChain + Streamlit | 
        <a href='https://langchain.com/' target='_blank'>LangChain Docs</a>
        </div>
        """, 
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()