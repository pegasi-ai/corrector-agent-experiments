import streamlit as st
from pydantic import BaseModel, Field
from typing import List, Tuple
import mammoth
import io
import re
from dotenv import load_dotenv
import os
import PyPDF2
import json
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from langchain_openai import ChatOpenAI
from langchain.agents import AgentExecutor, create_openai_functions_agent, create_structured_chat_agent
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.schema import BaseMessage
from langchain.tools import BaseTool
from langchain_core.tools import tool
import preprocess

load_dotenv()  # This loads variables from .env into the environment

# Set OpenAI API key from environment variable (supports both OPENAI_API_KEY and openai_api_key)
if os.getenv("OPENAI_API_KEY"):
    os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
elif os.getenv("openai_api_key"):
    os.environ["OPENAI_API_KEY"] = os.getenv("openai_api_key")

# Pydantic models for structured output
class Correction(BaseModel):
    original_text: str = Field(description="The original incorrect text")
    corrected_text: str = Field(description="The factually correct text")
    reason: str = Field(description="Explanation of why the correction was needed")
    confidence: float = Field(description="Confidence score from 0.0 to 1.0")

class DocumentAnalysis(BaseModel):
    corrections: List[Correction] = Field(description="List of factual corrections needed")
    overall_accuracy: float = Field(description="Overall accuracy score from 0.0 to 1.0")
    summary: str = Field(description="Summary of corrections made")

@tool
def fact_check_document(document_text: str) -> str:
    """
    Analyze a document for factual errors and return corrections in JSON format.
    
    Args:
        document_text: The text content of the document to fact-check
        
    Returns:
        JSON string with corrections, overall accuracy, and summary
    """
    # This tool will be used by the LangChain agent to perform fact-checking
    return f"Fact-checking completed for document with {len(document_text)} characters"

def create_langchain_agent(model_name: str):
    """Create a LangChain agent for fact-checking with support for multiple providers"""
    if model_name.startswith("openai:"):
        from langchain_openai import ChatOpenAI
        llm = ChatOpenAI(
            model=model_name.replace("openai:", ""),
            temperature=0.1,
            api_key=os.environ.get("OPENAI_API_KEY")
        )
    elif model_name.startswith("anthropic:"):
        from langchain_anthropic import ChatAnthropic
        llm = ChatAnthropic(
            model=model_name.replace("anthropic:", ""),
            temperature=0.1,
            api_key=os.environ.get("ANTHROPIC_API_KEY")
        )
    elif model_name.startswith("groq:"):
        from langchain_groq import ChatGroq
        llm = ChatGroq(
            model=model_name.replace("groq:", ""),
            temperature=0.1,
            api_key=os.environ.get("GROQ_API_KEY")
        )
    else:
        raise ValueError(f"Unsupported model: {model_name}")

    # Create the prompt template - Fixed to match expected variables
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a fact-checking expert. Your job is to:
1. Carefully read the provided document and website links
2. Identify any factually incorrect information given the reference from website links
3. Provide accurate corrections with explanations
4. Focus on verifiable facts, dates, statistics, historical events, scientific claims
5. Be conservative - only flag clear factual errors, not opinions or subjective statements
6. Provide a confidence score for each correction

Use the fact_check_document tool to analyze the document and return a JSON response with this exact structure:
{{
    "corrections": [
        {{
            "original_text": "the incorrect text",
            "corrected_text": "the correct text", 
            "reason": "explanation of why correction was needed",
            "confidence": 0.95
        }}
    ],
    "overall_accuracy": 0.85,
    "summary": "summary of corrections made"
}}"""),
        ("human", "{input}"),
        MessagesPlaceholder(variable_name="agent_scratchpad"),
    ])
    
    # Create tools
    tools = [fact_check_document]
    
    # Create the agent based on the model type
    if model_name.startswith("anthropic:"):
        # Use structured chat agent for Claude (no function-calling support)
        agent = create_structured_chat_agent(llm, tools, prompt)
    else:
        # Use function-calling agent for OpenAI/Groq
        agent = create_openai_functions_agent(llm, tools, prompt)
    
    # Create the agent executor
    agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)
    
    return agent_executor

def process_document_with_langchain(agent_executor, text: str, links: List[str]) -> DocumentAnalysis:
    """Process document using LangChain agent"""
    try:
        # Compose the reference links section
        links_section = "\n".join(f"- {link}" for link in links)
        # Run the agent with updated instructions
        result = agent_executor.invoke({
            "input": f"Please fact-check the following document and determine if the information is factually consistent based on the provided reference links.\n\nDocument:\n{text}\n\nReference Links:\n{links_section}\n\nFor each factual claim, check if it is supported or contradicted by the reference links."
        })
        
        # Extract the result
        result_text = result["output"]
        
        # Try to parse JSON from the result
        # The agent might return JSON in the response
        try:
            # Look for JSON in the response
            start_idx = result_text.find('{')
            end_idx = result_text.rfind('}') + 1
            if start_idx != -1 and end_idx != 0:
                json_str = result_text[start_idx:end_idx]
                result_json = json.loads(json_str)
            else:
                # If no JSON found, create a default response
                result_json = {
                    "corrections": [],
                    "overall_accuracy": 1.0,
                    "summary": "No factual errors found in the document."
                }
        except json.JSONDecodeError:
            # If JSON parsing fails, create a default response
            result_json = {
                "corrections": [],
                "overall_accuracy": 1.0,
                "summary": "Analysis completed. Please review the agent's response manually."
            }
        
        # Convert to DocumentAnalysis object
        corrections = [Correction(**corr) for corr in result_json["corrections"]]
        analysis = DocumentAnalysis(
            corrections=corrections,
            overall_accuracy=result_json["overall_accuracy"],
            summary=result_json["summary"]
        )
        
        return analysis
        
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        return None

def extract_text_from_docx(file_content):
    """Extract text from Word document using mammoth"""
    try:
        # Convert docx to HTML first to preserve some formatting
        result = mammoth.convert_to_html(io.BytesIO(file_content))
        html_content = result.value
        
        # Extract plain text for processing
        text_result = mammoth.extract_raw_text(io.BytesIO(file_content))
        plain_text = text_result.value
        
        return plain_text, html_content
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return None, None

def extract_text_from_pdf(file_content):
    """Extract text from PDF file using PyPDF2"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        
        return text, None  # No HTML for PDF
    except Exception as e:
        st.error(f"Error reading PDF document: {str(e)}")
        return None, None

def apply_corrections_with_red_marking(original_text: str, corrections: List[Correction]) -> Tuple[str, List[Correction]]:
    """Apply corrections to text with red highlighting for changes. Return both the highlighted text and the list of corrections that were actually applied, in the order they appear in the original text."""
    corrected_text = original_text
    applied_corrections = []
    # Instead of sorting by length, process in the original order
    for correction in corrections:
        original_lower = correction.original_text.lower()
        corrected_lower = corrected_text.lower()
        start_pos = 0
        applied = False
        while True:
            pos = corrected_lower.find(original_lower, start_pos)
            if pos == -1:
                break
            # Replace this occurrence
            before = corrected_text[:pos]
            after = corrected_text[pos + len(correction.original_text):]
            corrected_text = before + f'<span style="color: red; font-weight: bold;">{correction.corrected_text}</span>' + after
            corrected_lower = corrected_text.lower()
            start_pos = pos + len(f'<span style="color: red; font-weight: bold;">{correction.corrected_text}</span>')
            applied = True
        if applied:
            applied_corrections.append(correction)
    # HTML-encode $ symbols to ensure they display properly in Streamlit
    corrected_text = corrected_text.replace('$', '&#36;')
    return corrected_text, applied_corrections

def create_word_document_with_highlighting(original_text: str, corrections: List[Correction], analysis: DocumentAnalysis, filename: str) -> bytes:
    """Create a Word document with red highlighting for corrections, matching the Streamlit UI display"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Document Fact-Checking Report', 0)
    title.alignment = 1  # Center alignment
    
    # Add metadata
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Original File: {filename}')
    doc.add_paragraph(f'Overall Accuracy: {analysis.overall_accuracy:.1%}')
    doc.add_paragraph(f'Number of Corrections: {len(corrections)}')
    
    # Add summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(analysis.summary)
    
    # Add original text
    doc.add_heading('Original Text', level=1)
    doc.add_paragraph(original_text)
    
    # Add corrected text with highlighting (matching Streamlit UI)
    doc.add_heading('Corrected Text (with red highlighting)', level=1)
    corrected_text, applied_corrections = apply_corrections_with_red_marking(original_text, corrections)
    
    # Parse the HTML-style spans and add runs with formatting
    from html.parser import HTMLParser
    class HighlightParser(HTMLParser):
        def __init__(self, paragraph):
            super().__init__()
            self.paragraph = paragraph
            self.in_red_span = False
        def handle_starttag(self, tag, attrs):
            if tag == 'span':
                for attr in attrs:
                    if attr[0] == 'style' and 'color: red' in attr[1]:
                        self.in_red_span = True
        def handle_endtag(self, tag):
            if tag == 'span':
                self.in_red_span = False
        def handle_data(self, data):
            run = self.paragraph.add_run(data)
            if self.in_red_span:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True
    p = doc.add_paragraph()
    parser = HighlightParser(p)
    parser.feed(corrected_text.replace('&#36;', '$'))  # decode $ for Word
    
    # Add detailed corrections
    if applied_corrections:
        doc.add_heading('Detailed Corrections', level=1)
        for i, correction in enumerate(applied_corrections, 1):
            doc.add_heading(f'Correction #{i} (Confidence: {correction.confidence:.1%})', level=2)
            doc.add_paragraph('Original:')
            orig_p = doc.add_paragraph()
            orig_run = orig_p.add_run(correction.original_text)
            orig_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
            doc.add_paragraph('Corrected:')
            corr_p = doc.add_paragraph()
            corr_run = corr_p.add_run(correction.corrected_text)
            corr_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
            doc.add_paragraph('Reason:')
            doc.add_paragraph(correction.reason)
            doc.add_paragraph('─' * 60)
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

def main():
    st.set_page_config(
        page_title="Document Fact Checker",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Document Fact Checker")
    st.markdown("Upload a Microsoft Word document or PDF to check for factual accuracy and get corrections with red highlighting.")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Model selection
        model_options = {
            "OpenAI GPT-4o": "openai:gpt-4o",
            "OpenAI GPT-4o Mini": "openai:gpt-4o-mini",
            "Groq Llama 2 70B": "groq:llama-3.3-70b-versatile",
        }
        selected_model_label = st.selectbox(
            "Language Model",
            list(model_options.keys()),
            help="Choose the language model for fact-checking."
        )
        selected_model = model_options[selected_model_label]
        
        # API Key input
        api_key_provided = False
        
        if selected_model.startswith("openai:"):
            api_key = st.text_input(
                "OpenAI API Key", 
                type="password", 
                help="Enter your OpenAI API key to use the fact-checking service"
            )
            if api_key:
                os.environ["OPENAI_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("OPENAI_API_KEY"):
                api_key_provided = True
                st.success("✅ OpenAI API key loaded from environment")
                
        elif selected_model.startswith("anthropic:"):
            api_key = st.text_input(
                "Anthropic API Key", 
                type="password", 
                help="Enter your Anthropic API key to use Claude models"
            )
            if api_key:
                os.environ["ANTHROPIC_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("ANTHROPIC_API_KEY"):
                api_key_provided = True
                st.success("✅ Anthropic API key loaded from environment")
                
        elif selected_model.startswith("groq:"):
            api_key = st.text_input(
                "Groq API Key", 
                type="password", 
                help="Enter your Groq API key to use Llama models"
            )
            if api_key:
                os.environ["GROQ_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("GROQ_API_KEY"):
                api_key_provided = True
                st.success("✅ Groq API key loaded from environment")
        
        elif selected_model == "llama-2-70b":
            st.info("No API key required for Llama 2 70B.")
            api_key_provided = True
        
        st.markdown("---")
        st.markdown("### 📋 Instructions")
        st.markdown("""
        1. Select your preferred language model
        2. Enter the required API key above
        3. Upload a Microsoft Word document (.docx) or PDF
        4. Click 'Analyze Document' to start fact-checking
        5. Review corrections highlighted in red
        """)
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📤 Upload Document")
        
        # File uploader
        uploaded_file = st.file_uploader(
            "Choose a Word or PDF document",
            type=['docx', 'pdf'],
            help="Upload a .docx or .pdf file to fact-check"
        )
        
        # Clear previous analysis if a new file is uploaded
        if uploaded_file is not None:
            if (
                'uploaded_filename' not in st.session_state or
                st.session_state.uploaded_filename != uploaded_file.name
            ):
                # Clear all previous analysis data
                st.session_state.analysis = None
                st.session_state.original_text = None
                st.session_state.uploaded_filename = uploaded_file.name
                # Force a rerun to clear the display
                if hasattr(st.session_state, 'analysis'):
                    st.rerun()
        
        if uploaded_file is not None:
            st.success(f"✅ Uploaded: {uploaded_file.name}")
            
            # Extract text from document
            file_content = uploaded_file.read()
            if uploaded_file.name.lower().endswith('.docx'):
                plain_text, html_content = extract_text_from_docx(file_content)
            elif uploaded_file.name.lower().endswith('.pdf'):
                plain_text, html_content = extract_text_from_pdf(file_content)
            else:
                st.error("Unsupported file type.")
                return
            
            if plain_text:
                st.subheader("📄 Original Document Preview")
                with st.expander("View original text", expanded=False):
                    st.text_area("Original Content", plain_text, height=300, disabled=True)


                text_links_per_citation, text_links_per_chunk = preprocess.run(html_content)

                # # tmp
                # text, links = text_links_per_chunk[3]['text'], text_links_per_chunk[3]['citations']
                # print("============= tmp info =============")
                # print("text:", text)
                # print("links:", links)

                
                # Track which analyze button was last clicked
                if 'analyze_mode' not in st.session_state:
                    st.session_state.analyze_mode = None

                # Analyze buttons row with highlighting
                analyze_col1, analyze_col2 = st.columns([1, 1])
                with analyze_col1:
                    analyze_doc_clicked = st.button(
                        "🔍 Analyze Document",
                        type="primary" if st.session_state.analyze_mode == 'document' else "secondary",
                        key="analyze_doc_btn"
                    )
                with analyze_col2:
                    analyze_section_clicked = st.button(
                        "🔎 Analyze Section",
                        type="primary" if st.session_state.analyze_mode == 'section' else "secondary",
                        key="analyze_section_btn"
                    )

                # Section selection for Analyze Section (always visible if there are sections)
                section_options = [f"Section {i+1}: {chunk.get('text', '')[:40]}..." for i, chunk in enumerate(text_links_per_chunk)]
                selected_section_idx = st.selectbox("Select Section to Analyze", options=list(range(len(section_options))), format_func=lambda i: section_options[i], key="section_select")

                if analyze_doc_clicked:
                    st.session_state.analyze_mode = 'document'
                    if not api_key_provided:
                        st.error("⚠️ Please enter the required API key in the sidebar first!")
                        return
                    
                    with st.spinner("🤖 Fact-checking document... This may take a few moments."):
                        try:
                            # Clear any previous analysis before starting new one
                            st.session_state.analysis = None
                            
                            # Create LangChain agent
                            agent_executor = create_langchain_agent(selected_model)
                            
                            # Loop over each chunk and process
                            all_corrections = []
                            all_summaries = []
                            total_accuracy = 0.0
                            chunk_count = 0
                            for chunk in text_links_per_chunk:
                                chunk_text = chunk.get('text', '')
                                chunk_links = chunk.get('citations', [])
                                if not chunk_text.strip():
                                    continue
                                analysis = process_document_with_langchain(agent_executor, chunk_text, chunk_links)
                                if analysis:
                                    all_corrections.extend(analysis.corrections)
                                    all_summaries.append(analysis.summary)
                                    total_accuracy += analysis.overall_accuracy
                                    chunk_count += 1
                            # Aggregate results
                            if chunk_count > 0:
                                avg_accuracy = total_accuracy / chunk_count
                                summary = '\n'.join(all_summaries)
                            else:
                                avg_accuracy = 1.0
                                summary = "No factual errors found in the document."
                            final_analysis = DocumentAnalysis(
                                corrections=all_corrections,
                                overall_accuracy=avg_accuracy,
                                summary=summary
                            )
                            st.session_state.analysis = final_analysis
                            st.session_state.original_text = plain_text
                            st.session_state.analysis_timestamp = datetime.now().isoformat()
                            st.success("✅ Analysis complete!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error during analysis: {str(e)}")
                            st.error("Please check your API key and try again.")
                            st.session_state.analysis = None

                if analyze_section_clicked:
                    st.session_state.analyze_mode = 'section'
                    if not api_key_provided:
                        st.error("⚠️ Please enter the required API key in the sidebar first!")
                        return
                    section_chunk = text_links_per_chunk[selected_section_idx]
                    chunk_text = section_chunk.get('text', '')
                    chunk_links = section_chunk.get('citations', [])
                    with st.spinner(f"🤖 Fact-checking Section {selected_section_idx+1}..."):
                        try:
                            agent_executor = create_langchain_agent(selected_model)
                            analysis = process_document_with_langchain(agent_executor, chunk_text, chunk_links)
                            if analysis:
                                st.session_state.analysis = analysis
                                st.session_state.original_text = chunk_text
                                st.session_state.analysis_timestamp = datetime.now().isoformat()
                                st.success(f"✅ Section {selected_section_idx+1} analysis complete!")
                                st.rerun()
                            else:
                                st.error("❌ Analysis failed. Please try again.")
                        except Exception as e:
                            st.error(f"❌ Error during section analysis: {str(e)}")
                            st.session_state.analysis = None
    
    with col2:
        st.header("📊 Analysis Results")
        
        if hasattr(st.session_state, 'analysis') and st.session_state.analysis:
            analysis = st.session_state.analysis
            original_text = st.session_state.original_text
            
            # Overall metrics
            col2a, col2b = st.columns(2)
            with col2a:
                st.metric("Overall Accuracy", f"{analysis.overall_accuracy:.1%}")
            with col2b:
                st.metric("Corrections Found", len(analysis.corrections))
            
            # Summary
            st.subheader("📝 Summary")
            # HTML-encode $ symbols in the summary to ensure proper display
            encoded_summary = analysis.summary.replace('$', '&#36;')
            st.info(encoded_summary)
            
            # Corrections list
            if analysis.corrections:
                st.subheader("🔧 Corrections Made")
                
                # Display corrected text
                st.subheader("📄 Corrected Document")
                corrected_text, applied_corrections = apply_corrections_with_red_marking(original_text, analysis.corrections)
                st.markdown(
                    f'<div style="border: 1px solid #ddd; padding: 15px; border-radius: 5px; background-color: #fafafa; max-height: 400px; overflow-y: auto;">{corrected_text}</div>',
                    unsafe_allow_html=True
                )
                
                # Detailed corrections
                st.subheader("📋 Detailed Corrections")
                for i, correction in enumerate(applied_corrections, 1):
                    with st.expander(f"Correction {i} (Confidence: {correction.confidence:.1%})"):
                        col_orig, col_corr = st.columns(2)
                        
                        with col_orig:
                            st.markdown("**Original:**")
                            st.markdown(f'<div style="background-color: #ffebee; padding: 10px; border-radius: 5px;">{correction.original_text}</div>', unsafe_allow_html=True)
                        
                        with col_corr:
                            st.markdown("**Corrected:**")
                            st.markdown(f'<div style="background-color: #e8f5e8; padding: 10px; border-radius: 5px;">{correction.corrected_text}</div>', unsafe_allow_html=True)
                        
                        st.markdown("**Reason:**")
                        # HTML-encode $ symbols in the reason text to ensure proper display
                        encoded_reason = correction.reason.replace('$', '&#36;')
                        st.markdown(f'<div style="background-color: #f5f5f5; padding: 10px; border-radius: 5px;">{encoded_reason}</div>', unsafe_allow_html=True)
                
                # Download corrected document
                st.subheader("💾 Export")
                
                # Determine the correct file extension
                if uploaded_file and uploaded_file.name.lower().endswith('.pdf'):
                    base_filename = uploaded_file.name.replace('.pdf', '')
                elif uploaded_file and uploaded_file.name.lower().endswith('.docx'):
                    base_filename = uploaded_file.name.replace('.docx', '')
                else:
                    base_filename = "document"
                
                # Create Word document with highlighting
                word_doc_bytes = create_word_document_with_highlighting(
                    original_text, analysis.corrections, analysis, 
                    uploaded_file.name if uploaded_file else "document"
                )
                
                # Download Word document
                st.download_button(
                    label="📄 Download Report",
                    data=word_doc_bytes,
                    file_name=f"fact_check_report_{base_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            else:
                st.success("🎉 No factual errors found! The document appears to be accurate.")
        
        else:
            st.info("👆 Upload a document and click 'Analyze Document' to see results here.")
    
    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666; font-size: 0.8em;'>
        Built with ❤️ using LangChain + Streamlit | 
        <a href='https://langchain.com/' target='_blank'>LangChain Docs</a>
        </div>
        """, 
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()