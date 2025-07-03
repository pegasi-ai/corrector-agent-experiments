from docx import Document
import mammoth
import textwrap
from html import unescape
from typing import Dict, Tuple, List, Any
import re
from langchain.text_splitter import RecursiveCharacterTextSplitter
import mistune


# load file
def docx_to_html(path: str) -> str:
    """Convert a .docx into HTML, preserving tables and links."""
    with open(path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        # result.value is the HTML string
        # result.messages contains any warnings
        return result.value
    
def split_plaintext_into_sections(plain_text: str, chunk_size: int = 1000, chunk_overlap: int = 100) -> list:
    """
    Split plain text into sections using LangChain's RecursiveCharacterTextSplitter.
    Returns a list of section strings.
    """
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_text(plain_text)

# split data
def split_html_into_paragraphs(html):
    
    # First, extract tables and replace them with placeholders to avoid interference
    table_placeholders = []
    table_counter = 0
    
    def replace_table_with_placeholder(match):
        nonlocal table_counter
        placeholder = f"__TABLE_PLACEHOLDER_{table_counter}__"
        table_placeholders.append((placeholder, match.group(0)))
        table_counter += 1
        return placeholder
    
    # Replace tables with placeholders (non-greedy to handle nested tables)
    html_without_tables = re.sub(r'<table.*?>.*?</table>', replace_table_with_placeholder, html, flags=re.DOTALL | re.IGNORECASE)
    
    # Use regex to find all <p>...</p> blocks in the HTML without tables
    paragraphs = re.findall(r'<p.*?>.*?</p>', html_without_tables, re.DOTALL | re.IGNORECASE)
    
    # Filter paragraphs to only include those with meaningful content (words or numbers)
    meaningful_paragraphs = []
    for paragraph in paragraphs:
        # Remove HTML tags to get plain text
        plain_text = re.sub(r'<[^>]+>', '', paragraph)
        # Decode HTML entities
        plain_text = unescape(plain_text)
        # Remove extra whitespace
        plain_text = re.sub(r'\s+', ' ', plain_text).strip()
        
        # Check if paragraph contains words or numbers
        if re.search(r'[a-zA-Z0-9]', plain_text):
            meaningful_paragraphs.append(paragraph)
    
    # Restore tables from placeholders and add them as paragraphs
    for placeholder, table_html in table_placeholders:
        # Replace the placeholder in the original HTML with the actual table
        meaningful_paragraphs.append(table_html)
    
    return meaningful_paragraphs

# parse file
def extract_text_and_citations_per_citation_href(html: str) -> Dict[str, str]:
    """
    Finds all <a href="...">...</a> links in the HTML string,
    and for each returns the text immediately before the link (stripped of any tags).
    """
    # Regex to find <a ... href="URL">label</a>
    link_re = re.compile(
        r'<a\s+[^>]*href="(?P<href>[^"]+)"[^>]*>.*?</a>',
        flags=re.IGNORECASE | re.DOTALL
    )
    
    results: Dict[str, str] = {}
    last_end = 0
    
    for m in link_re.finditer(html):
        url = m.group('href')
        # Segment from end of last link (or start) up to this link
        pre_segment = html[last_end : m.start()]
        # Remove any HTML tags from that segment and collapse whitespace
        text_before = re.sub(r'<[^>]+>', '', pre_segment).strip()
        results[text_before] = url
        last_end = m.end()
    
    return results

def extract_text_and_citations_per_section_href(text: str) -> dict:
    """
    Same functionality but returns a dictionary instead of tuple.
    Returns:
    dict: {"text": clean_text, "citations": citations}
    """
    # 1. Extract all citation URLs
    citations = re.findall(
        r'<a\s+[^>]*href="([^"]+)"',
        text,
        flags=re.IGNORECASE
    )
    # 2. Remove all <a>…</a> segments entirely (so link text doesn't remain)
    html_no_links = re.sub(
        r'<a\s+[^>]*>.*?</a>',
        '',
        text,
        flags=re.IGNORECASE | re.DOTALL
    )
    # 3. Strip all remaining HTML tags
    text_only = re.sub(r'<[^>]+>', ' ', html_no_links)
    # 4. Decode HTML entities (e.g. &amp; → &)
    text_only = unescape(text_only)
    # 5. Collapse multiple whitespace into single spaces
    clean_text = re.sub(r'\s+', ' ', text_only).strip()
    
    return {
        "text": clean_text,
        "citations": citations
    }

def extract_text_and_citations_per_citation_source(html: str) -> Dict[str, str]:
    """
    From an HTML string with markdown-style [Source](URL) citations,
    return a dict mapping each citation context to its URL.
    Each context is the clean text chunk immediately preceding that citation.
    """
    # 1. Decode HTML entities
    decoded = unescape(html)

    # 2. Strip HTML tags
    text = re.sub(r'<[^>]+>', ' ', decoded)

    # 3. Normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()

    # 4. Find all [Source](URL)
    link_pattern = re.compile(r'\[Source\]\((https?://[^\)]+)\)', flags=re.IGNORECASE)

    results: Dict[str, str] = {}
    last_end = 0

    for m in link_pattern.finditer(text):
        url = m.group(1)
        # the text between the end of the previous link (or start) and this one
        context = text[last_end:m.start()].strip()
        results[context] = url
        last_end = m.end()

    return results

# case I
def run_href(text: str):
 
    # II. split file
    sections = split_html_into_paragraphs(text)

    # III. parse file
    text_links_per_citation = []
    text_links_per_section = []
    for section in sections:
        dic = extract_text_and_citations_per_citation_href(section)
        if not dic:
            continue

        text_links_per_citation.append(dic)

        dic = extract_text_and_citations_per_section_href(section)
        text_links_per_section.append(dic)

    return text_links_per_citation, text_links_per_section


# case II
def run_source(text: str):
 
    # II. split file
    sections = split_html_into_paragraphs(text)

    # III. parse file
    text_links_per_citation = []
    for section in sections:
        dic = extract_text_and_citations_per_citation_source(section)
        if not dic:
            continue
        text_links_per_citation.append(dic)


    return text_links_per_citation

def extract_text_images_tables_from_md(md_content: str) -> tuple:
    """
    Extract plain text (excluding figures, tables, images), HTML, tables, and figures from a Markdown string.
    Returns (plain_text, html_content, images, tables, figures)
    - plain_text: all text content, excluding figures, tables, and images
    - html_content: HTML rendering of the markdown
    - images: list of image URLs/paths (from markdown and HTML)
    - tables: list of HTML tables (from both Markdown and raw HTML)
    - figures: list of <figure>...</figure> blocks (as strings)
    """
    # Remove <figure>...</figure> blocks
    no_figures = re.sub(r'<figure[\s\S]*?>[\s\S]*?</figure>', '', md_content, flags=re.IGNORECASE)
    # Remove Markdown tables (lines with | and at least one header separator)
    no_tables = re.sub(r'(^\s*\|.*\|\s*$\n?)+', '', no_figures, flags=re.MULTILINE)
    # Remove HTML tables
    no_tables = re.sub(r'<table[\s\S]*?>[\s\S]*?</table>', '', no_tables, flags=re.IGNORECASE)
    # Remove Markdown images ![alt](url)
    no_images = re.sub(r'!\[[^\]]*\]\(([^)]+)\)', '', no_tables)
    # Remove HTML images <img ...>
    no_images = re.sub(r'<img [^>]*src=["\"][^"\"]+["\"][^>]*>', '', no_images, flags=re.IGNORECASE)

    # Render HTML
    markdown = mistune.create_markdown(renderer=mistune.HTMLRenderer())
    html_content = markdown(md_content)

    # Use mistune's AST to extract tables and text
    class Collector(mistune.HTMLRenderer):
        def __init__(self):
            super().__init__()
            self.tables = []
            self.text_chunks = []
        def table(self, header, body):
            table_html = f"<table>{header}{body}</table>"
            self.tables.append(table_html)
            return table_html
        def text(self, text):
            self.text_chunks.append(text)
            return super().text(text)
    collector = Collector()
    mistune.create_markdown(renderer=collector)(no_images)
    plain_text = " ".join(collector.text_chunks).strip()

    # Extract <figure>...</figure> blocks using regex (non-greedy)
    figure_pattern = re.compile(r'<figure[\s\S]*?>[\s\S]*?</figure>', re.IGNORECASE)
    figures = figure_pattern.findall(md_content)

    # Extract images from markdown ![alt](url) and HTML <img ... src="...">
    md_img_pattern = re.compile(r'!\[[^\]]*\]\(([^)]+)\)')
    html_img_pattern = re.compile(r'<img [^>]*src=["\\\']([^"\\\']+)["\\\']', re.IGNORECASE)
    images = md_img_pattern.findall(md_content) + html_img_pattern.findall(md_content)

    # Extract raw HTML tables
    html_table_pattern = re.compile(r'<table[\s\S]*?>[\s\S]*?</table>', re.IGNORECASE)
    html_tables = html_table_pattern.findall(md_content)

    # Combine Markdown and HTML tables, avoiding duplicates
    all_tables = collector.tables.copy()
    for tbl in html_tables:
        if tbl not in all_tables:
            all_tables.append(tbl)

    return plain_text, html_content, images, all_tables, figures



    