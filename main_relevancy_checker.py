import streamlit as st
from pydantic import BaseModel, Field
from typing import List
import mammoth
import io
import re
from dotenv import load_dotenv
import os
import PyPDF2
import json
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from langchain_openai import ChatOpenAI
from langchain.agents import AgentExecutor, create_openai_functions_agent, create_structured_chat_agent
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.schema import BaseMessage
from langchain.tools import BaseTool
from langchain_core.tools import tool
import preprocess

load_dotenv()  # This loads variables from .env into the environment

# Set OpenAI API key from environment variable (supports both OPENAI_API_KEY and openai_api_key)
if os.getenv("OPENAI_API_KEY"):
    os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
elif os.getenv("openai_api_key"):
    os.environ["OPENAI_API_KEY"] = os.getenv("openai_api_key")

# Pydantic models for structured output
class Relevancy(BaseModel):
    original_text: str = Field(description="The original text")
    website_link: str = Field(description="The associated website link to original text")
    relevant: bool = Field(description="A boolean of Yes or No to indicate if the link is relevant to the original text or not")
    reason: str = Field(description="Explanation of why website link is irrelevant to original text")
    confidence: float = Field(description="Confidence score from 0.0 to 1.0")
    

class DocumentAnalysis(BaseModel):
    relevancies: List[Relevancy] = Field(description="List of original texts and the associated website links")

@tool
def relevancy_check_document(document_text: str) -> str:
    """
    Analyze a document for irrelevant errors and return irrelevancies in JSON format.
    Args:
        document_text: The text content of the document to relevancy-check
    Returns:
        JSON string with relevancies 
    """
    return f"Relevancy-checking completed for document with {len(document_text)} characters"

def create_langchain_agent(model_name: str):
    """Create a LangChain agent for relevancy-checking with support for multiple providers"""
    if model_name.startswith("openai:"):
        from langchain_openai import ChatOpenAI
        llm = ChatOpenAI(
            model=model_name.replace("openai:", ""),
            temperature=0.1,
            api_key=os.environ.get("OPENAI_API_KEY")
        )
    elif model_name.startswith("anthropic:"):
        from langchain_anthropic import ChatAnthropic
        llm = ChatAnthropic(
            model=model_name.replace("anthropic:", ""),
            temperature=0.1,
            api_key=os.environ.get("ANTHROPIC_API_KEY")
        )
    elif model_name.startswith("groq:"):
        from langchain_groq import ChatGroq
        llm = ChatGroq(
            model=model_name.replace("groq:", ""),
            temperature=0.1,
            api_key=os.environ.get("GROQ_API_KEY")
        )
    else:
        raise ValueError(f"Unsupported model: {model_name}")

    # Create the prompt template - Fixed to match expected variables
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a relevancy-checking expert. Your job is to:\n1. Carefully read the provided document and website link\n2. Identify if the website link provides RELEVANT and SUFFICIENT information to the document\n3. Provide the final response explanations\n5. Be conservative - only flag irrelevent link to the document, not opinions or subjective statements\n6. Provide a confidence score for each statement\n\nUse the relevancy_check_document tool to analyze the document and return a JSON response with this exact structure:\n{{\n    \"relevancies\": [\n        {{\n            \"original_text\": \"the original document associated with the website link\",\n            \"website_link\": \"the website link associated with the original text\"\n        \"relevant\": \"a boolean of Yes or No to indicate if the website link is relevant to the original text or not\"\n        \"reason\": \"explanation of why the website link is irrelevant to the original text\",\n            \"confidence\": 0.95,\n            }}\n    ]\n}}"""),    
        ("human", "{input}"),
        MessagesPlaceholder(variable_name="agent_scratchpad"),
    ])
    
    # Create tools
    tools = [relevancy_check_document]
    
    # Create the agent based on the model type
    if model_name.startswith("anthropic:"):
        # Use structured chat agent for Claude (no function-calling support)
        agent = create_structured_chat_agent(llm, tools, prompt)
    else:
        # Use function-calling agent for OpenAI/Groq
        agent = create_openai_functions_agent(llm, tools, prompt)
    
    # Create the agent executor
    agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)
    
    return agent_executor

def process_document_with_langchain(agent_executor, text: str, link: str) -> DocumentAnalysis:
    """Process document using LangChain agent for fact-checking."""
    try:
        result = agent_executor.invoke({
            "input": f"Given a text and a reference link, determine if the provided reference link provides relevant and sufficient information for the text. \n\Text:\n{text}\n\nReference Link:\n{link}"
        })
        result_text = result["output"]
        try:
            start_idx = result_text.find('{')
            end_idx = result_text.rfind('}') + 1
            if start_idx != -1 and end_idx != 0:
                json_str = result_text[start_idx:end_idx]
                result_json = json.loads(json_str)
            else:
                result_json = {
                    "relevancies": [],
                }
        except json.JSONDecodeError:
            result_json = {
                "relevancies": [],
            }
        relevancies = [Relevancy(**corr) for corr in result_json["relevancies"]]
        analysis = DocumentAnalysis(
            relevancies=relevancies,
        )
        return analysis
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        return None

def extract_text_from_docx(file_content):
    """Extract text from Word document using mammoth"""
    try:
        # Convert docx to HTML first to preserve some formatting
        result = mammoth.convert_to_html(io.BytesIO(file_content))
        html_content = result.value
        
        # Extract plain text for processing
        text_result = mammoth.extract_raw_text(io.BytesIO(file_content))
        plain_text = text_result.value
        
        return plain_text, html_content
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return None, None

def extract_text_from_pdf(file_content):
    """Extract text from PDF file using PyPDF2"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        
        return text, None  # No HTML for PDF
    except Exception as e:
        st.error(f"Error reading PDF document: {str(e)}")
        return None, None

# Add function to create Word document with relevancy report
def create_relevancy_report_word_document(relevancies: List[Relevancy], filename: str) -> bytes:
    """Create a Word document with relevancy report data"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Document Relevancy Report', 0)
    title.alignment = 1  # Center alignment
    
    # Add metadata
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Original File: {filename}')
    doc.add_paragraph(f'Total Links Analyzed: {len(relevancies)}')
    
    # Add relevancy data
    if relevancies:
        doc.add_heading('Relevancy Analysis Results', level=1)
        for i, relevancy in enumerate(relevancies, 1):
            doc.add_heading(f'Link Analysis #{i} (Confidence: {relevancy.confidence:.1%})', level=2)
            
            doc.add_paragraph('Original Text:')
            doc.add_paragraph(relevancy.original_text)
            
            doc.add_paragraph('Website Link:')
            doc.add_paragraph(relevancy.website_link)
            
            doc.add_paragraph('Relevancy:')
            doc.add_paragraph(str(relevancy.relevant))
            
            doc.add_paragraph('Reason:')
            doc.add_paragraph(relevancy.reason)
            
            doc.add_paragraph('─' * 60)
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

def main():
    st.set_page_config(
        page_title="Document Relevancy Checker",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Document Relevancy Checker")
    st.markdown("Upload a Microsoft Word document or PDF to check for relevancy.")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Model selection
        model_options = {
            "OpenAI GPT-4o": "openai:gpt-4o",
            # "OpenAI GPT-4o Mini": "openai:gpt-4o-mini",
            "OpenAI GPT-4.1 Mini": "openai:gpt-4.1-mini",
            # "Groq Llama-3.3 70B": "groq:llama-3.3-70b-versatile",
        }
        selected_model_label = st.selectbox(
            "Language Model",
            list(model_options.keys()),
            help="Choose the language model for relevancy-checking."
        )
        selected_model = model_options[selected_model_label]
        
        # API Key input
        api_key_provided = False
        
        if selected_model.startswith("openai:"):
            api_key = st.text_input(
                "OpenAI API Key", 
                type="password", 
                help="Enter your OpenAI API key to use the relevancy-checking service"
            )
            if api_key:
                os.environ["OPENAI_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("OPENAI_API_KEY"):
                api_key_provided = True
                st.success("✅ OpenAI API key loaded from environment")
                
        elif selected_model.startswith("anthropic:"):
            api_key = st.text_input(
                "Anthropic API Key", 
                type="password", 
                help="Enter your Anthropic API key to use Claude models"
            )
            if api_key:
                os.environ["ANTHROPIC_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("ANTHROPIC_API_KEY"):
                api_key_provided = True
                st.success("✅ Anthropic API key loaded from environment")
                
        elif selected_model.startswith("groq:"):
            api_key = st.text_input(
                "Groq API Key", 
                type="password", 
                help="Enter your Groq API key to use Llama models"
            )
            if api_key:
                os.environ["GROQ_API_KEY"] = api_key
                api_key_provided = True
            elif os.getenv("GROQ_API_KEY"):
                api_key_provided = True
                st.success("✅ Groq API key loaded from environment")
        
        elif selected_model == "llama-2-70b":
            st.info("No API key required for Llama 2 70B.")
            api_key_provided = True
        
        st.markdown("---")
        st.markdown("### 📋 Instructions")
        st.markdown("""
        1. Select your preferred language model
        2. Enter the required API key above
        3. Upload a Microsoft Word document (.docx) or PDF
        4. Click 'Analyze Document' to start relevancy-checking for the whole document or 'Analyze Section' for a selected section
        """)
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📤 Upload Document")
        
        # File uploader
        uploaded_file = st.file_uploader(
            "Choose a Word or PDF document",
            type=['docx', 'pdf'],
            help="Upload a .docx or .pdf file to relevancy-check"
        )
        
        # Clear previous analysis if a new file is uploaded
        if uploaded_file is not None:
            if (
                'uploaded_filename' not in st.session_state or
                st.session_state.uploaded_filename != uploaded_file.name
            ):
                # Clear all previous analysis data
                st.session_state.analysis = None
                st.session_state.original_text = None
                st.session_state.uploaded_filename = uploaded_file.name
                # Force a rerun to clear the display
                if hasattr(st.session_state, 'analysis'):
                    st.rerun()
        
        if uploaded_file is not None:
            st.success(f"✅ Uploaded: {uploaded_file.name}")
            
            # Extract text from document
            file_content = uploaded_file.read()
            if uploaded_file.name.lower().endswith('.docx'):
                plain_text, html_content = extract_text_from_docx(file_content)
            elif uploaded_file.name.lower().endswith('.pdf'):
                plain_text, html_content = extract_text_from_pdf(file_content)
            else:
                st.error("Unsupported file type.")
                return
            
            if plain_text:
                st.subheader("📄 Original Document Preview")
                with st.expander("View original text", expanded=False):
                    st.text_area("Original Content", plain_text, height=300, disabled=True)


                text_links_per_citation = preprocess.run_source(html_content)

                # Track which analyze button was last clicked
                if 'analyze_mode' not in st.session_state:
                    st.session_state.analyze_mode = None

                # Analyze buttons row with highlighting
                analyze_col1, analyze_col2 = st.columns([1, 1])
                with analyze_col1:
                    analyze_doc_clicked = st.button(
                        "🔍 Analyze Document",
                        type="primary" if st.session_state.analyze_mode == 'document' else "secondary",
                        key="analyze_doc_btn"
                    )
                with analyze_col2:
                    analyze_section_clicked = st.button(
                        "🔎 Analyze Section",
                        type="primary" if st.session_state.analyze_mode == 'section' else "secondary",
                        key="analyze_section_btn"
                    )

                # Section selection for Analyze Section (always visible if there are sections)
                section_options = [f"Section {i+1}: {', '.join(list(chunk.keys())[:1])[:40]}..." for i, chunk in enumerate(text_links_per_citation)]
                selected_section_idx = st.selectbox("Select Section to Analyze", options=list(range(len(section_options))), format_func=lambda i: section_options[i], key="section_select")

                if analyze_doc_clicked:
                    st.session_state.analyze_mode = 'document'
                    if not api_key_provided:
                        st.error("⚠️ Please enter the required API key in the sidebar first!")
                        return
                    with st.spinner("🤖 Relevancy-checking document... This may take a few moments."):
                        try:
                            st.session_state.analysis = None
                            agent_executor = create_langchain_agent(selected_model)
                            all_relevancies = []
                            section_count = 0
                            for section_dict in text_links_per_citation:
                                for text, link in section_dict.items():
                                    if not text.strip():
                                        continue
                                    analysis = process_document_with_langchain(agent_executor, text, [link])
                                    if analysis:
                                        all_relevancies.extend(analysis.relevancies)
                                        section_count += 1
                        
                            final_analysis = DocumentAnalysis(
                                relevancies=all_relevancies,
                            )
                            st.session_state.analysis = final_analysis
                            st.session_state.original_text = plain_text
                            st.session_state.analysis_timestamp = datetime.now().isoformat()
                            st.success("✅ Analysis complete!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error during analysis: {str(e)}")
                            st.error("Please check your API key and try again.")
                            st.session_state.analysis = None

                if analyze_section_clicked:
                    st.session_state.analyze_mode = 'section'
                    if not api_key_provided:
                        st.error("⚠️ Please enter the required API key in the sidebar first!")
                        return
                    if selected_section_idx is None or not text_links_per_citation or selected_section_idx >= len(text_links_per_citation):
                        st.error("No section selected or section index is invalid.")
                        return
                    section_dict = text_links_per_citation[selected_section_idx]
                   
                    with st.spinner(f"🤖 Relevancy-checking Section {selected_section_idx+1}..."):
                        try:
                            agent_executor = create_langchain_agent(selected_model)
                            all_relevancies = []
                            section_count = 0
                            for text, link in section_dict.items():
                                if not text.strip():
                                    continue

                                analysis = process_document_with_langchain(agent_executor, text, link)
                                if analysis:
                                    all_relevancies.extend(analysis.relevancies)
                                    section_count += 1
                
                            final_analysis = DocumentAnalysis(
                                relevancies=all_relevancies,
                            )
                            st.session_state.analysis = final_analysis
                            st.session_state.original_text = '\n\n'.join(section_dict.keys())
                            st.session_state.analysis_timestamp = datetime.now().isoformat()
                            st.success(f"✅ Section {selected_section_idx+1} analysis complete!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error during section analysis: {str(e)}")
                            st.session_state.analysis = None
    
    with col2:
        st.header("📊 Analysis Results")
        
        if hasattr(st.session_state, 'analysis') and st.session_state.analysis:
            analysis = st.session_state.analysis
            original_text = st.session_state.original_text
            
            # Overall metrics
            col2a, col2b = st.columns(2)
            with col2a:
                st.metric("Number of Links", len(analysis.relevancies))
            with col2b:
                irrelevant_count = sum(1 for relevancy in analysis.relevancies if not relevancy.relevant or relevancy.relevant == "No")
                st.metric("Number of Irrelevant Links", irrelevant_count)
            
            # Relevancy list
            if analysis.relevancies:
                
                # Detailed relevancies
                st.subheader("📋 Detailed Explanation")
                for i, relevancy in enumerate(analysis.relevancies, 1):
                    with st.expander(f"Link {i} (Confidence: {relevancy.confidence:.1%})"):
                        st.markdown("**Original:**")
                        st.markdown(f'<div style="background-color: #ffebee; padding: 10px; border-radius: 5px;">{relevancy.original_text}</div>', unsafe_allow_html=True)
                        st.markdown("**Website Link:**")
                        st.markdown(f'<div style="background-color: #e3f2fd; padding: 10px; border-radius: 5px; word-break: break-all;">{relevancy.website_link}</div>', unsafe_allow_html=True)
                        st.markdown("**Relevancy:**")
                        st.markdown(f'<div style="background-color: #e3f2fd; padding: 10px; border-radius: 5px; word-break: break-all;">{relevancy.relevant}</div>', unsafe_allow_html=True)
                        st.markdown("**Reason:**")
                        encoded_reason = relevancy.reason.replace('$', '&#36;')
                        st.markdown(f'<div style="background-color: #f5f5f5; padding: 10px; border-radius: 5px;">{encoded_reason}</div>', unsafe_allow_html=True)
            
                # Add download button for Word report
                st.subheader("💾 Export Report")
                if uploaded_file:
                    base_filename = uploaded_file.name.replace('.pdf', '').replace('.docx', '')
                else:
                    base_filename = "document"
                
                word_doc_bytes = create_relevancy_report_word_document(analysis.relevancies, uploaded_file.name if uploaded_file else "document")
                st.download_button(
                    label="📄 Download Relevancy Report",
                    data=word_doc_bytes,
                    file_name=f"relevancy_report_{base_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.success("🎉 No irrelevant errors found! The links appear to be relevant.")
        
        else:
            st.info("👆 Upload a document and click 'Analyze Document' to see results here.")
    
    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666; font-size: 0.8em;'>
        Built with ❤️ using LangChain + Streamlit | 
        <a href='https://langchain.com/' target='_blank'>LangChain Docs</a>
        </div>
        """, 
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()